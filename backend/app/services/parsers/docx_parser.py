import logging
import re
from pathlib import Path
from typing import BinaryIO

from docx import Document as DocxDocument

from app.models.document_object import DocumentObject, PageContent
from app.utils.pii_redaction import redact_pii

logger = logging.getLogger(__name__)

HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Title"}
SECTION_KEYWORDS = re.compile(
    r"\b(Termination|Liability|Payment|Confidentiality|Arbitration)\b",
    re.IGNORECASE,
)


def parse_docx(source: str | Path | BinaryIO) -> DocumentObject:
    if isinstance(source, (str, Path)):
        doc = DocxDocument(str(source))
        metadata = {"format": "docx", "path": str(source)}
    else:
        doc = DocxDocument(source)
        metadata = {"format": "docx"}

    paragraphs: list[str] = []
    sections: list[str] = []
    title = None

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        text = redact_pii(text)
        style = para.style.name if para.style else ""
        if style in HEADING_STYLES or SECTION_KEYWORDS.search(text):
            if len(text) < 120:
                sections.append(text)
                if title is None:
                    title = text
        paragraphs.append(text)

    full_text = "\n".join(paragraphs)
    # DOCX has no native pages — approximate ~500 words per page
    words = full_text.split()
    page_size = 500
    pages: list[PageContent] = []
    for i in range(0, max(len(words), 1), page_size):
        chunk_words = words[i : i + page_size]
        page_text = " ".join(chunk_words)
        page_num = (i // page_size) + 1
        page_sections = [s for s in sections if s.lower() in page_text.lower()]
        pages.append(PageContent(page_number=page_num, text=page_text, sections=page_sections))

    if not pages and full_text:
        pages = [PageContent(page_number=1, text=full_text, sections=sections)]

    core_props = doc.core_properties
    if core_props.title:
        title = core_props.title
    metadata.update(
        {
            "author": core_props.author,
            "subject": core_props.subject,
        }
    )

    return DocumentObject(
        text=full_text,
        pages=pages,
        metadata=metadata,
        title=title or (full_text[:80] if full_text else "Untitled"),
    )
